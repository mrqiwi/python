#!/usr/bin/env python3

from docx import Document
from docx.shared import Inches
import os

filename = 'demo.docx'
dirname = 'output'

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered list', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

m = document.add_paragraph('hello ')
m.add_run('how are you').bold = True


# document.add_picture('Christmas_Holidays_Gifts_428526.jpg', width=Inches(1.25))
document.add_picture('Christmas_Holidays_Gifts_428526.jpg', width=Inches(5.25), height=Inches(1))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'

for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.add_heading('Начинаем вторую страницу', level=1)
document.add_paragraph('	Жил-был хоббит в норе')

if not os.path.exists(dirname):
	os.mkdir(dirname)

os.chdir(dirname)

document.save(filename)
